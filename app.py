# app.py
#!/usr/bin/env python3
"""
Streamlit: PDF → OCR (Mistral on Azure) → Markdown → DOCX
- Same UX as your sample app (auth, S3 silent upload, page wallet, admin)
- Uses your Mistral OCR endpoint (base64 data URL upload)
- Pandoc path renders LaTeX math ($...$, $$...$$) and tables
- Optional: embed OCR-returned images; optional region crops (via JSON)
- Fixes "Could not fetch resource" by stripping inline image refs from text and
  embedding local image files with Pandoc --resource-path
"""

import io
import os
import json
import re
import time
import base64
import hashlib
import shutil
import tempfile
from typing import Optional, Any, Dict, List, Tuple
from datetime import datetime, timezone
from pathlib import Path

import requests
import streamlit as st
import boto3
from botocore.exceptions import ClientError
from docx import Document
from docx.shared import Pt, Inches

# ── Pandoc / pypandoc (optional but recommended for math/tables) ──────────────
HAVE_PYPANDOC = True
try:
    import pypandoc  # type: ignore
except Exception:
    HAVE_PYPANDOC = False

def ensure_pandoc_available() -> bool:
    """Return True if system pandoc is available, or can be downloaded via pypandoc."""
    if shutil.which("pandoc"):
        return True
    if HAVE_PYPANDOC:
        try:
            pypandoc.download_pandoc()
            return True
        except Exception:
            return False
    return False

# =========================
# PAGE SETUP
# =========================
st.set_page_config(page_title="PDF → DOCX Suvichaars", page_icon="📄", layout="wide")
st.title("📄 PDF to DOCX with Suvichaar Doc AI")
st.caption(
    "Upload a PDF → OCR extracts text → Download a .docx • "
    "Each PDF page deducts 1 page from your balance • Default balance: 10,000 pages (admin can top-up)"
)

# =========================
# PAGE BALANCE MODEL
# =========================
DEFAULT_START_PAGES = 10_000  # per-user starting page allowance

# =========================
# SECRETS / CONFIG (NO HARDCODED DEFAULTS)
# =========================
def get_secret(key: str, default: Optional[str] = None) -> Optional[str]:
    try:
        return st.secrets[key]  # type: ignore[attr-defined]
    except Exception:
        return os.environ.get(key, default)

# ---- SuvichaarOCR (Mistral on Azure AI) ----
MISTRAL_OCR_ENDPOINT = get_secret("MISTRAL_OCR_ENDPOINT")
MISTRAL_API_KEY      = get_secret("MISTRAL_API_KEY")
MISTRAL_MODEL        = get_secret("MISTRAL_MODEL", "mistral-document-ai-2505")

# --- AWS / S3 (silent uploads) ---
AWS_REGION            = get_secret("AWS_REGION", "ap-south-1")
AWS_ACCESS_KEY_ID     = get_secret("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = get_secret("AWS_SECRET_ACCESS_KEY")
AWS_SESSION_TOKEN     = get_secret("AWS_SESSION_TOKEN")  # optional

S3_BUCKET = get_secret("S3_BUCKET", "suvichaarapp")
S3_PREFIX = (get_secret("S3_PREFIX", "media/pdf2docx") or "media/pdf2docx").lstrip("/")

# --- Admin bootstrap ---
ADMIN_EMAIL    = get_secret("ADMIN_EMAIL")
ADMIN_PASSWORD = get_secret("ADMIN_PASSWORD")

# --- Admin Panel PIN (6 digits) ---
ADMIN_PANEL_PIN = (get_secret("ADMIN_PANEL_PIN") or "").strip()
if not re.fullmatch(r"\d{6}", ADMIN_PANEL_PIN):
    st.error("ADMIN_PANEL_PIN missing/invalid. Set a 6-digit PIN in .streamlit/secrets.toml or .env and restart.")
    st.stop()

# Quick guards for required secrets
required = {
    "MISTRAL_OCR_ENDPOINT": MISTRAL_OCR_ENDPOINT,
    "MISTRAL_API_KEY": MISTRAL_API_KEY,
    "ADMIN_EMAIL": ADMIN_EMAIL,
    "ADMIN_PASSWORD": ADMIN_PASSWORD,
}
missing = [k for k, v in required.items() if not v]
if missing:
    st.error(f"Missing required config: {', '.join(missing)}. Add them to .streamlit/secrets.toml or .env.")
    st.stop()

# =========================
# USERS STORE (auth + page wallet)
# =========================
USERS_STORE_PATH = Path(
    os.getenv("USERS_STORE_PATH", "")
    or (Path(os.getenv("USERS_STORE_DIR", Path(tempfile.gettempdir()) / "suvichaar_pdfdoc")) / "users_store.json")
)

def _ensure_store_parent() -> None:
    USERS_STORE_PATH.parent.mkdir(parents=True, exist_ok=True)

APP_SALT = b"SuvichaarDI_v1"  # app-level salt for PBKDF2

def _hash_pw(password: str, salt: bytes) -> str:
    h = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, 120_000)
    return base64.b64encode(h).decode("utf-8")

def _set_pw(password: str) -> str:
    return _hash_pw(password, APP_SALT)

DEFAULT_USERS_DB = {"users": {}}  # email -> record

def load_users() -> Dict[str, Any]:
    _ensure_store_parent()
    if USERS_STORE_PATH.exists():
        try:
            with open(USERS_STORE_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return DEFAULT_USERS_DB.copy()
    return DEFAULT_USERS_DB.copy()

def save_users(data: Dict[str, Any]) -> None:
    _ensure_store_parent()
    tmp_path = USERS_STORE_PATH.with_suffix(USERS_STORE_PATH.suffix + ".tmp")
    with open(tmp_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp_path, USERS_STORE_PATH)

# =========================
# SESSION BOOTSTRAP
# =========================
if "users_db" not in st.session_state:
    st.session_state.users_db = load_users()
if "current_user" not in st.session_state:
    st.session_state.current_user = None
if "auth_view" not in st.session_state:
    st.session_state.auth_view = "login"
if "admin_panel_unlocked" not in st.session_state:
    st.session_state.admin_panel_unlocked = False

# First-run: ensure admin exists
def _migrate_to_pages_model(rec: Dict[str, Any]) -> Dict[str, Any]:
    if "start_pages" not in rec and "start_credits" in rec:
        rec["start_pages"] = int(rec.get("start_credits", DEFAULT_START_PAGES))
    if "remaining_pages" not in rec and "credits" in rec:
        rec["remaining_pages"] = int(rec.get("credits", rec.get("start_pages", DEFAULT_START_PAGES)))
    if "start_pages" not in rec:
        rec["start_pages"] = DEFAULT_START_PAGES
    if "remaining_pages" not in rec:
        rec["remaining_pages"] = rec.get("start_pages", DEFAULT_START_PAGES)
    return rec

if not ADMIN_EMAIL or not ADMIN_PASSWORD:
    st.error("ADMIN_EMAIL / ADMIN_PASSWORD not set in config.")
    st.stop()

if ADMIN_EMAIL not in st.session_state.users_db["users"]:
    st.session_state.users_db["users"][ADMIN_EMAIL] = _migrate_to_pages_model({
        "email": ADMIN_EMAIL,
        "name": "Admin",
        "tenant_id": "default-tenant",
        "profile_id": "admin-profile",
        "password_hash": _set_pw(ADMIN_PASSWORD),
        "force_pw_change": False,
        "is_admin": True,
        "start_pages": DEFAULT_START_PAGES,
        "remaining_pages": DEFAULT_START_PAGES,
        "ledger": [],
        "charged_docs": {},
        "last_txn": None,
        "last_s3_keys": [],
    })
    save_users(st.session_state.users_db)
else:
    admin_rec = st.session_state.users_db["users"][ADMIN_EMAIL]
    st.session_state.users_db["users"][ADMIN_EMAIL] = _migrate_to_pages_model(admin_rec)
    save_users(st.session_state.users_db)

# =========================
# S3 HELPERS (silent uploads) — no ACLs
# =========================
def _sanitize_filename(name: str) -> str:
    base = name.strip().replace(" ", "_")
    return re.sub(r"[^A-Za-z0-9._-]+", "", base) or "file"

def _have_static_creds() -> bool:
    return bool(AWS_ACCESS_KEY_ID and AWS_SECRET_ACCESS_KEY)

@st.cache_resource(show_spinner=False)
def _get_s3_client(_region: str, _ak: Optional[str], _sk: Optional[str], _st: Optional[str]):
    if _ak and _sk:
        session = boto3.session.Session(
            aws_access_key_id=_ak,
            aws_secret_access_key=_sk,
            aws_session_token=_st,
            region_name=_region,
        )
    else:
        session = boto3.session.Session(region_name=_region)
    return session.client("s3")

def _build_object_key(prefix: str, kind: str, tenant_id: str, email: str, fid: str, filename: str, ext: str) -> str:
    safe = _sanitize_filename(filename.rsplit(".", 1)[0])
    today = datetime.now().strftime("%Y/%m/%d")
    email_key = email.replace("@", "_")
    return f"{(prefix or 'media/pdf2docx').rstrip('/')}/{kind}/{tenant_id}/{email_key}/{today}/{fid[:12]}-{safe}.{ext.lstrip('.')}"

def _put_bytes_to_s3(key: str, data: bytes, content_type: str) -> None:
    extra = {"ContentType": content_type}
    client = _get_s3_client(AWS_REGION, AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, AWS_SESSION_TOKEN)
    client.put_object(Bucket=S3_BUCKET, Key=key, Body=data, **extra)

def silent_upload_pdf(fid: str, filename: str, pdf_bytes: bytes, tenant_id: str, email: str):
    try:
        key = _build_object_key(S3_PREFIX, "uploads", tenant_id, email, fid, filename, "pdf")
        _put_bytes_to_s3(key, pdf_bytes, "application/pdf")
        rec = get_user_rec()
        (rec.setdefault("last_s3_keys", [])).append({"type": "pdf", "key": key, "ts": datetime.now().isoformat()})
        save_user_rec(rec)
    except Exception:
        pass  # silent by design

def silent_upload_docx(fid: str, filename: str, docx_bytes: bytes, tenant_id: str, email: str):
    try:
        key = _build_object_key(S3_PREFIX, "outputs", tenant_id, email, fid, filename, "docx")
        _put_bytes_to_s3(key, docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        rec = get_user_rec()
        (rec.setdefault("last_s3_keys", [])).append({"type": "docx", "key": key, "ts": datetime.now().isoformat()})
        save_user_rec(rec)
    except Exception:
        pass  # silent by design

def run_s3_health_check():
    try:
        client = _get_s3_client(AWS_REGION, AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, AWS_SESSION_TOKEN)
        source = "static keys" if _have_static_creds() else "instance/role"
        st.info(f"Using AWS credentials from: {source}")

        loc = client.get_bucket_location(Bucket=S3_B
UCKET).get("LocationConstraint") or "us-east-1"
        st.success(f"S3 bucket location: {loc}")

        test_key = f"{(S3_PREFIX or 'media/pdf2docx').rstrip('/')}/healthcheck/{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S')}.txt"
        client.put_object(Bucket=S3_BUCKET, Key=test_key, Body=b"ok", ContentType="text/plain")
        st.success(f"Put OK → {test_key}")

        obj = client.get_object(Bucket=S3_BUCKET, Key=test_key)
        data = obj["Body"].read()
        st.success(f"Get OK ({len(data)} bytes)")

        client.delete_object(Bucket=S3_BUCKET, Key=test_key)
        st.success("Delete OK")
    except ClientError as e:
        err = e.response.get("Error", {})
        st.error(f"S3 ClientError: {err.get('Code')} — {err.get('Message')}")
    except Exception as e:
        st.error(f"S3 health check failed: {e}")

# =========================
# PER-USER HELPERS
# =========================
def get_user_rec() -> Dict[str, Any]:
    rec = st.session_state.current_user
    rec = _migrate_to_pages_model(rec)
    st.session_state.current_user = rec
    save_user_rec(rec)
    return rec

def save_user_rec(rec: Dict[str, Any]) -> None:
    db = st.session_state.users_db
    db["users"][rec["email"]] = rec
    save_users(db)
    st.session_state.current_user = rec

def file_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()

def debit_user_pages(rec: Dict[str, Any], fid: str, pages: int, filename: str) -> int:
    pages = max(1, int(pages))
    charged_docs = rec.setdefault("charged_docs", {})
    if fid in charged_docs:
        return 0  # already debited for this file hash

    remaining = int(rec.get("remaining_pages", 0))
    if remaining < pages:
        raise RuntimeError(f"Insufficient page balance: need {pages}, have {remaining}.")

    rec["remaining_pages"] = remaining - pages
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    txn = {"file": filename, "pages": pages, "ts": ts}
    rec["last_txn"] = txn
    (rec.setdefault("ledger", [])).append({"file": filename, "pages": pages, "ts": ts})
    charged_docs[fid] = txn
    save_user_rec(rec)
    return pages

# =========================
# AUTH UI
# =========================
def ui_login():
    st.subheader("Login")
    email = st.text_input("Email", key="auth_email")
    pw = st.text_input("Password", type="password", key="auth_pw")

    if st.button("Sign in", key="auth_signin_btn", use_container_width=True):
        rec = st.session_state.users_db["users"].get(email)
        if not rec:
            st.error("Invalid email or password.")
            return

        if rec.get("force_pw_change"):
            temp_hash = rec.get("temp_pw_hash")
            if temp_hash and _set_pw(pw) == temp_hash:
                st.session_state.auth_view = "reset"
                st.session_state.reset_email_prefill = email
                st.info("First-time login detected. Please set a new password to continue.")
                return
            else:
                st.error("This account requires a password reset. Use your temporary password to proceed.")
                return

        if _set_pw(pw) == rec.get("password_hash"):
            st.session_state.current_user = rec
            st.success(f"Welcome {rec.get('name') or rec['email']}!")
        else:
            st.error("Invalid email or password.")

def ui_reset_password():
    st.subheader("Reset Password")
    st.caption("Use the temporary password you received from Admin once, then set a new password.")
    email = st.text_input("Email", value=st.session_state.get("reset_email_prefill",""), key="reset_email")
    temp_pw = st.text_input("Temporary Password", type="password", key="reset_temp")
    new_pw = st.text_input("New Password", type="password", key="reset_new")
    new_pw2 = st.text_input("Re-enter New Password", type="password", key="reset_new2")
    if st.button("Reset Password", key="reset_btn", use_container_width=True):
        rec = st.session_state.users_db["users"].get(email)
        if not rec:
            st.error("Account not found.")
            return
        if _set_pw(temp_pw) != rec.get("temp_pw_hash"):
            st.error("Temporary password incorrect.")
            return
        if not new_pw or new_pw != new_pw2:
            st.error("New passwords do not match.")
            return
        rec["password_hash"] = _set_pw(new_pw)
        rec["force_pw_change"] = False
        rec.pop("temp_pw_hash", None)
        save_user_rec(rec)
        st.success("Password updated. Please login.")
        st.session_state.auth_view = "login"

with st.sidebar:
    st.markdown("### Navigation")
    nav_choice = st.radio(
        label="",
        options=["Login", "Reset Password"],
        index=0 if st.session_state.auth_view == "login" else 1,
        key="auth_nav",
    )
    st.session_state.auth_view = "login" if nav_choice == "Login" else "reset"

# Gate until login succeeds
if st.session_state.current_user is None:
    if st.session_state.auth_view == "login":
        ui_login()
    else:
        ui_reset_password()
    st.stop()

# =========================
# SIDEBAR: PROFILE + PAGE BALANCE + ADMIN
# =========================
with st.sidebar:
    u = get_user_rec()

    # Profile card
    st.markdown(
        f"""
        <div style="background:#111827;border:1px solid #374151;border-radius:12px;padding:12px;margin-bottom:10px;">
          <div style="font-size:16px;font-weight:600;">👤 Profile</div>
          <div style="font-size:13px;opacity:.9;margin-top:6px;">
            <div><b>Email:</b> {u['email']}</div>
            <div><b>Tenant ID:</b> {u.get('tenant_id','-')}</div>
            <div><b>Profile ID:</b> {u.get('profile_id','-')}</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Page balance meter
    st.subheader("📄 Page Balance")
    start_cap = max(int(u.get("start_pages", DEFAULT_START_PAGES)), int(u.get("remaining_pages", 0)))
    pct = max(0.0, min(float(u.get("remaining_pages", 0)) / float(start_cap or 1), 1.0))
    st.progress(pct, text=f"Pages remaining: {int(u.get('remaining_pages', 0))}")
    st.caption("Each extracted page deducts 1 page from your balance.")

    # Last transaction (optional)
    txn = u.get("last_txn")
    if txn:
        st.markdown(
            f"""
            <div style="background:#f5f8ff;padding:12px;border-radius:10px;border:1px solid #d1e3ff;margin-top:12px;">
              <div style="font-weight:600;color:#1f4396;margin-bottom:6px;">🧾 Last Transaction</div>
              <div style="font-size:13px;line-height:1.4;">
                <div><b>File:</b> {txn['file']}</div>
                <div><b>Pages Debited:</b> {txn['pages']}</div>
                <div style="color:#666;"><b>Time:</b> {txn.get('ts','')}</div>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    # Admin panel (only if user is admin) + requires 6-digit PIN
    if u.get("is_admin"):
        with st.expander("🔐 Admin Panel", expanded=False):
            if not st.session_state.admin_panel_unlocked:
                st.info("Enter the 6-digit Admin PIN to unlock the panel.")
                pin_in = st.text_input("Admin PIN (6 digits)", type="password", max_chars=6, key="admin_pin_input")
                if st.button("Unlock Admin Panel", key="admin_pin_btn"):
                    if re.fullmatch(r"\d{6}", str(pin_in or "")) and str(pin_in) == ADMIN_PANEL_PIN:
                        st.session_state.admin_panel_unlocked = True
                        st.success("Admin Panel unlocked.")
                    else:
                        st.error("Invalid PIN. Please try again.")
            else:
                if st.button("🔒 Lock Admin Panel", key="admin_pin_lock"):
                    st.session_state.admin_panel_unlocked = False

                st.markdown("**Create / Edit User**")
                a_email = st.text_input("User Email", key="a_email")
                a_name = st.text_input("Name", key="a_name")
                a_tenant = st.text_input("Tenant ID", key="a_tenant")
                a_profile = st.text_input("Profile ID", key="a_profile")
                a_start_pages = st.number_input("Start Pages", min_value=0, value=DEFAULT_START_PAGES, step=100, key="a_start_pages")
                a_pages_now = st.number_input("Current Pages (remaining)", min_value=0, value=DEFAULT_START_PAGES, step=100, key="a_pages_now")
                a_temp_pw = st.text_input("Temporary Password (for new/reset)", type="password", key="a_temp")

                if st.button("Save User", key="a_save"):
                    if not a_email:
                        st.error("Email required.")
                    else:
                        db = st.session_state.users_db
                        rec = db["users"].get(a_email, {})
                        rec = _migrate_to_pages_model(rec or {})
                        rec.update({
                            "email": a_email,
                            "name": a_name or rec.get("name") or "",
                            "tenant_id": a_tenant or rec.get("tenant_id") or "",
                            "profile_id": a_profile or rec.get("profile_id") or "",
                            "is_admin": rec.get("is_admin", False),
                            "start_pages": int(a_start_pages),
                            "remaining_pages": int(a_pages_now),
                            "ledger": rec.get("ledger", []),
                            "charged_docs": rec.get("charged_docs", {}),
                            "last_txn": rec.get("last_txn", None),
                            "last_s3_keys": rec.get("last_s3_keys", []),
                        })
                        if a_temp_pw:
                            rec["temp_pw_hash"] = _set_pw(a_temp_pw)
                            rec["force_pw_change"] = True
                        else:
                            rec["force_pw_change"] = rec.get("force_pw_change", False)

                        db["users"][a_email] = rec
                        save_users(db)
                        st.success("User saved / updated.")

                st.markdown("---")
                st.markdown("**Top-up Pages**")
                top_email = st.text_input("Email to top-up", key="top_email")
                add_pages = st.number_input("Pages to add", min_value=1, value=100, step=50, key="top_pages")
                if st.button("Top-up", key="top_btn"):
                    db = st.session_state.users_db
                    rec = db["users"].get(top_email)
                    if not rec:
                        st.error("User not found.")
                    else:
                        rec = _migrate_to_pages_model(rec)
                        rec["remaining_pages"] = int(rec.get("remaining_pages", 0)) + int(add_pages)
                        save_users(db)
                        st.success(f"Added {add_pages} pages to {top_email}.")

                st.markdown("---")
                st.markdown("**Grant/Revoke Admin**")
                adm_email = st.text_input("Email", key="adm_email")
                make_admin = st.checkbox("Is Admin?", value=False, key="adm_flag")
                if st.button("Update Admin Flag", key="adm_btn"):
                    db = st.session_state.users_db
                    rec = db["users"].get(adm_email)
                    if not rec:
                        st.error("User not found.")
                    else:
                        rec["is_admin"] = bool(make_admin)
                        save_users(db)
                        st.success("Updated.")

                st.markdown("---")
                st.markdown("**Set Tenant ID / Profile ID**")
                users_map = st.session_state.users_db.get("users", {})
                user_emails = sorted(users_map.keys())
                sel_email = st.selectbox("Select user", options=user_emails, key="tenant_profile_sel_email")

                if sel_email:
                    target = _migrate_to_pages_model(users_map.get(sel_email, {}))
                    cur_tenant  = target.get("tenant_id", "")
                    cur_profile = target.get("profile_id", "")
                    new_tenant  = st.text_input("Tenant ID",  value=cur_tenant,  key="tenant_profile_new_tenant")
                    new_profile = st.text_input("Profile ID", value=cur_profile, key="tenant_profile_new_profile")

                    if st.button("Save Tenant/Profile", key="tenant_profile_save_btn"):
                        target["tenant_id"]  = (new_tenant or "").strip()
                        target["profile_id"] = (new_profile or "").strip()
                        db = st.session_state.users_db
                        db["users"][sel_email] = target
                        save_users(db)
                        if st.session_state.current_user["email"] == sel_email:
                            st.session_state.current_user = target
                        st.success(f"Updated tenant/profile for {sel_email}.")

                st.markdown("---")
                st.markdown("**S3 Health Check**")
                if st.button("Run S3 Health Check", key="s3_health_btn"):
                    run_s3_health_check()

# =========================
# SETTINGS + DEBUG
# =========================
with st.expander("⚙️ Settings", expanded=False):
    add_page_breaks   = st.checkbox("Insert page breaks between PDF pages", value=True, key="opt_page_breaks")
    include_images    = st.checkbox("Include OCR-detected images (if present)", value=True, key="opt_include_images")
    image_max_width   = st.number_input("Image width (inches)", min_value=1.0, max_value=8.5, value=6.5, step=0.5, key="opt_img_width")

    prefer_pandoc     = st.checkbox("Prefer Pandoc for DOCX (better math/tables)", value=True, key="opt_prefer_pandoc")
    force_pandoc      = st.checkbox("Force Pandoc even if no math/tables detected", value=False, key="opt_force_pandoc")
    offer_md_download = st.checkbox("Offer Markdown download (debug)", value=False, key="opt_offer_md")
    show_raw_json     = st.checkbox("Show raw OCR JSON (debug)", value=False, key="opt_raw_json")

    st.markdown("---")
    st.caption("Optional: upload a regions.json to crop figures/equations/tables from the PDF and embed them:")
    crop_json = st.file_uploader("regions.json (optional)", type=["json"], key="regions_json_uploader")

# =========================
# OCR HELPERS (Mistral) — base64-only + markdown-aware
# =========================
def bytes_to_data_url(mime: str, data: bytes) -> str:
    b64 = base64.b64encode(data).decode("utf-8")
    return f"data:{mime};base64,{b64}"

def _post_ocr(payload: Dict[str, Any]) -> Dict[str, Any]:
    headers = {
        "Authorization": f"Bearer {MISTRAL_API_KEY}",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }
    resp = requests.post(MISTRAL_OCR_ENDPOINT, headers=headers, json=payload, timeout=600)
    try:
        resp.raise_for_status()
    except requests.HTTPError:
        raise requests.HTTPError(f"{resp.status_code} {resp.reason}: {resp.text[:1500]}")
    try:
        return resp.json()
    except Exception:
        return {"raw": resp.text}

def call_mistral_ocr_pdf(pdf_bytes: bytes, want_images: bool) -> Dict[str, Any]:
    payload = {
        "model": MISTRAL_MODEL,
        "document": {
            "type": "document_url",
            "document_url": bytes_to_data_url("application/pdf", pdf_bytes),
        },
        "include_image_base64": bool(want_images)
    }
    return _post_ocr(payload)

# ---- markdown-first extraction + unwrapping + image collection
def _unwrap_container(obj: Dict[str, Any]) -> Dict[str, Any]:
    node = obj
    for k in ("output", "response", "result", "data", "ocr", "document"):
        if isinstance(node, dict) and isinstance(node.get(k), dict):
            node = node[k]
    return node

def _s(x): return (x or "").strip()

def _extract_from_page_dict(p: Dict[str, Any]) -> str:
    md = p.get("markdown")
    if isinstance(md, str) and md.strip():
        return md.strip()

    if isinstance(p.get("lines"), list):
        parts = []
        for ln in p["lines"]:
            if isinstance(ln, dict):
                t = _s(ln.get("content") or ln.get("text"))
                if t: parts.append(t)
        if parts: return "\n".join(parts)

    if isinstance(p.get("paragraphs"), list):
        parts = []
        for para in p["paragraphs"]:
            if isinstance(para, dict):
                t = _s(para.get("content") or para.get("text"))
                if t: parts.append(t)
        if parts: return "\n".join(parts)

    for key in ("blocks", "items", "elements", "regions"):
        arr = p.get(key)
        if isinstance(arr, list) and arr:
            parts = []
            for it in arr:
                if isinstance(it, dict):
                    t = _s(it.get("text") or it.get("content") or it.get("value"))
                    if t: parts.append(t)
            if parts: return "\n".join(parts)

    t = _s(p.get("content") or p.get("text") or p.get("full_text") or p.get("raw_text"))
    return t or ""

def extract_pages_texts_and_images(ocr_json: Dict[str, Any]) -> Tuple[List[str], Dict[int, List[Path]]]:
    """Return (pages_texts, images_by_page). Images saved into a temp 'assets' dir."""
    container = _unwrap_container(ocr_json)
    pages = container.get("pages")

    texts: List[str] = []
    images_by_page: Dict[int, List[Path]] = {}

    # Create a temp assets dir per run
    if "assets_dir" not in st.session_state:
        st.session_state.assets_dir = Path(tempfile.mkdtemp(prefix="pdf2docx_assets_"))
    assets_dir = st.session_state.assets_dir
    assets_dir.mkdir(parents=True, exist_ok=True)

    def _save_base64_image_unknown(input_str: str, stem: str) -> Optional[Path]:
        try:
            m = re.match(r'^data:(?P<mime>[^;]+);base64,(?P<b64>.+)$', input_str.strip(), re.IGNORECASE)
            b64 = m.group("b64") if m else input_str.strip()
            raw = base64.b64decode(b64, validate=False)
            from PIL import Image
            with Image.open(io.BytesIO(raw)) as im:
                fmt = (im.format or "").upper()
                if fmt == "PNG" and im.mode == "RGB":
                    out = assets_dir / f"{stem}.png"
                    im.save(out); return out
                im = im.convert("RGB")
                out = assets_dir / f"{stem}.jpg"
                im.save(out, quality=92, optimize=True); return out
        except Exception:
            return None

    if isinstance(pages, list) and pages:
        for i, p in enumerate(pages, start=1):
            texts.append(_extract_from_page_dict(p if isinstance(p, dict) else {}))

            # Collect base64 images (if present)
            if isinstance(p, dict):
                ims = p.get("images")
                if isinstance(ims, list) and ims:
                    for j, im in enumerate(ims, start=1):
                        raw = None
                        if isinstance(im, dict):
                            for k in ("base64","image_base64","data","imageData","content"):
                                v = im.get(k)
                                if isinstance(v, str) and len(v) > 50:
                                    raw = v; break
                            if not raw:
                                for k in ("data_url","url","image"):
                                    v = im.get(k)
                                    if isinstance(v, str) and v.lower().startswith("data:"):
                                        raw = v; break
                        if raw:
                            saved = _save_base64_image_unknown(raw, f"p{i}_{j}_ocrimg")
                            if saved:
                                images_by_page.setdefault(i, []).append(saved)

        if any(x.strip() for x in texts):
            return texts, images_by_page

    # Fallback to flat string fields
    for k in ("markdown", "full_text", "content", "text", "raw_text"):
        if isinstance(container.get(k), str) and container[k].strip():
            return [container[k].strip()], images_by_page

    # Ensure DOCX isn't empty
    return [json.dumps(ocr_json, ensure_ascii=False)], images_by_page

# =========================
# Markdown + DOCX builders
# =========================
def clean_markdown(md: str) -> str:
    md = md.replace("\r\n", "\n").replace("\r", "\n")
    return md

_MATH_OR_TABLE_RE = re.compile(
    r"(\$\$.*?\$\$|\$[^$\n]+\$|\\\(|\\\)|\\\[|\\\]|\\begin\{(equation|align|eqnarray|gather|aligned)\}|(^\s*\|.*\|\s*$\n^\s*\|?\s*[-:]+\s*(\|[-:]+\s*)+$))",
    re.MULTILINE | re.DOTALL
)

def detect_math_or_tables(pages_text: List[str]) -> bool:
    joined = "\n\n".join(pages_text)
    return bool(_MATH_OR_TABLE_RE.search(joined))

# Strip inline images so Pandoc doesn't try to fetch non-existent urls/ids
_IMG_MD_RE   = re.compile(r'!\[[^\]]*\]\([^)]+\)', re.IGNORECASE)  # ![alt](src)
_IMG_HTML_RE = re.compile(r'<img\b[^>]*>', re.IGNORECASE)          # <img ...>

def strip_inline_images(md: str) -> str:
    md = _IMG_MD_RE.sub('', md)
    md = _IMG_HTML_RE.sub('', md)
    return md

def md_image(path: Path, width_in: float) -> str:
    p = Path(path).resolve().as_posix()
    return f'![]({p}){{width={width_in}in}}'

def build_markdown(pages_text: List[str],
                   images_by_page: Dict[int, List[Path]],
                   crops_by_page: Dict[int, List[Path]],
                   insert_page_breaks: bool,
                   image_max_width_in: float) -> str:
    parts: List[str] = []
    for i, txt in enumerate(pages_text, start=1):
        text_clean = strip_inline_images(_s(txt))
        parts.append(f"\n\n## Page {i}\n\n{text_clean}\n")
        for im in images_by_page.get(i, []):
            parts.append("\n" + md_image(im, image_max_width_in) + "\n")
        for im in crops_by_page.get(i, []):
            parts.append("\n" + md_image(im, image_max_width_in) + "\n")
        if insert_page_breaks and i < len(pages_text):
            parts.append("\n\\newpage\n")
    return clean_markdown("".join(parts)).strip() + "\n"

def build_docx_with_pandoc(md_text: str, resource_dirs: Optional[List[Path]] = None) -> bytes:
    """Convert Markdown→DOCX using Pandoc (with math/tables support)."""
    if not HAVE_PYPANDOC:
        raise RuntimeError("pypandoc not installed")
    if not ensure_pandoc_available():
        raise RuntimeError("Pandoc not available and auto-download failed")

    with tempfile.TemporaryDirectory() as td:
        md_path   = Path(td) / "doc.md"
        docx_path = Path(td) / "out.docx"
        md_path.write_text(md_text, encoding="utf-8")

        extra_args = ["--standalone"]
        if resource_dirs:
            search_path = os.pathsep.join(str(Path(p).resolve()) for p in resource_dirs)
            extra_args.append(f"--resource-path={search_path}")

        pypandoc.convert_file(
            str(md_path),
            to="docx",
            format="gfm+tex_math_dollars+pipe_tables",
            outputfile=str(docx_path),
            extra_args=extra_args,
        )
        return docx_path.read_bytes()

def build_docx_with_python_docx(pages_text: List[str], insert_page_breaks: bool) -> bytes:
    """Plain paragraphs in DOCX (no native equations)."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if len(pages_text) > 1:
        for i, txt in enumerate(pages_text, start=1):
            doc.add_heading(f"Page {i}", level=2)
            for para in (txt or "").splitlines():
                if para.strip():
                    doc.add_paragraph(para)
            if insert_page_breaks and i < len(pages_text):
                doc.add_page_break()
    else:
        for para in (pages_text[0] or "").splitlines():
            if para.strip():
                doc.add_paragraph(para)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# =========================
# Optional: region crops (figures/equations/tables)
# =========================
from PIL import Image
import fitz  # PyMuPDF

def render_pdf_page_to_image(pdf_path: Path, page_num: int, dpi: int = 300):
    doc = fitz.open(pdf_path.as_posix())
    try:
        page = doc[page_num-1]
        zoom = dpi / 72.0
        pix  = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        img  = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        return img, (page.rect.width, page.rect.height)
    finally:
        doc.close()

def clamp(v, lo, hi): return max(lo, min(hi, v))

def bbox_to_pixels(b, img_w, img_h, pts_wh=None, coord_type="norm", origin="top-left"):
    x0,y0,x1,y1 = b
    if coord_type == "norm":
        if origin == "bottom-left":
            y0, y1 = 1 - y0, 1 - y1; y0, y1 = y1, y0
        X0,Y0,X1,Y1 = int(round(x0*img_w)), int(round(y0*img_h)), int(round(x1*img_w)), int(round(y1*img_h))
    elif coord_type == "pixel":
        if origin == "bottom-left":
            y0, y1 = img_h - y0, img_h - y1; y0, y1 = y1, y0
        X0,Y0,X1,Y1 = int(round(x0)), int(round(y0)), int(round(x1)), int(round(y1))
    elif coord_type == "pdf_points":
        if not pts_wh: raise ValueError("pdf_points requires page size in points")
        pts_w, pts_h = pts_wh; sx, sy = img_w/pts_w, img_h/pts_h
        if origin == "bottom-left":
            y0, y1 = pts_h - y0, pts_h - y1; y0, y1 = y1, y0
        X0,Y0,X1,Y1 = int(round(x0*sx)), int(round(y0*sy)), int(round(x1*sx)), int(round(y1*sy))
    else:
        raise ValueError("coord_type must be norm|pixel|pdf_points")
    x0p,x1p = sorted([X0,X1]); y0p,y1p = sorted([Y0,Y1])
    x0p = clamp(x0p,0,img_w-1); x1p = clamp(x1p,1,img_w)
    y0p = clamp(y0p,0,img_h-1); y1p = clamp(y1p,1,img_h)
    return x0p,y0p,x1p,y1p

def crop_and_save(img: Image.Image, bbox_px, padding, assets_dir: Path, stem: str) -> Path:
    x0,y0,x1,y1 = bbox_px
    if padding:
        x0 = max(0, x0-padding); y0 = max(0, y0-padding)
        x1 = min(img.width, x1+padding); y1 = min(img.height, y1+padding)
    crop = img.crop((x0,y0,x1,y1)).convert("RGB")
    assets_dir.mkdir(parents=True, exist_ok=True)
    out_path = assets_dir / f"{stem}.jpg"
    crop.save(out_path, quality=92, optimize=True)
    return out_path

def crops_from_regions(pdf_path: Path, regions_json: Dict[str, Any], dpi: int, assets_dir: Path) -> Dict[int, List[Path]]:
    pages_cfg = regions_json.get("pages", {})
    out: Dict[int, List[Path]] = {}
    for p_str, regs in pages_cfg.items():
        try:
            pnum = int(p_str)
        except:
            continue
        if not isinstance(regs, list):
            continue
        page_img, pts_wh = render_pdf_page_to_image(pdf_path, pnum, dpi=dpi)
        for k, r in enumerate(regs, start=1):
            coords     = r["coords"]
            coord_type = r.get("coord_type","norm")
            origin     = r.get("origin","top-left")
            padding    = int(r.get("padding", 8))
            label      = r.get("label", f"crop{k}")
            bbox_px = bbox_to_pixels(tuple(coords), page_img.width, page_img.height,
                                     pts_wh if coord_type=="pdf_points" else None,
                                     coord_type=coord_type, origin=origin)
            saved = crop_and_save(page_img, bbox_px, padding, assets_dir, f"p{pnum}_{k}_{label}")
            out.setdefault(pnum, []).append(saved)
    return out

# =========================
# MAIN FLOW
# =========================
uploaded = st.file_uploader("Upload a PDF", type=["pdf"], accept_multiple_files=False, key="pdf_uploader_main")

if uploaded is not None:
    if not uploaded.name.lower().endswith(".pdf"):
        st.error("Please upload a PDF file.")
    else:
        pdf_bytes = uploaded.read()
        if not pdf_bytes:
            st.error("Uploaded file is empty. Please re-upload the PDF.")
            st.stop()

        fid = file_hash(pdf_bytes)

        # Silent S3 upload (PDF) before analysis
        u = get_user_rec()
        silent_upload_pdf(fid, uploaded.name, pdf_bytes, u.get("tenant_id", "default"), u["email"])

        with st.spinner("Analyzing with Suvichaar Doc AI (OCR)..."):
            try:
                ocr_json = call_mistral_ocr_pdf(pdf_bytes, want_images=st.session_state.get("opt_include_images", True))
            except requests.HTTPError as e:
                st.error(f"OCR failed: {e}")
                st.stop()
            except Exception as e:
                st.error(f"OCR request failed: {e}")
                st.stop()

        if st.session_state.get("opt_raw_json"):
            with st.expander("Raw OCR JSON", expanded=False):
                st.json(ocr_json)

        # Extract page texts + collect OCR images (if any)
        pages_text, ocr_images_by_page = extract_pages_texts_and_images(ocr_json)
        pages_count = max(1, len(pages_text))
        st.success(f"Extracted text from **{pages_count} page(s)**.")

        # Quick first-page preview (helps sanity check)
        if pages_text and any(pages_text):
            st.code((pages_text[0] or "")[:1000], language="markdown")

        # Page debit (once per user+file hash)
        try:
            debited = debit_user_pages(u, fid, pages_count, uploaded.name)
            if debited > 0:
                st.toast(f"Debited {debited} page(s) from your balance.", icon="✅")
        except RuntimeError as e:
            st.error(str(e))
            st.stop()

        # Optional: load regions.json for crops
        crops_by_page: Dict[int, List[Path]] = {}
        if crop_json is not None:
            try:
                cfg = json.load(io.StringIO(crop_json.getvalue().decode("utf-8")))
                # Prepare session assets dir
                if "assets_dir" not in st.session_state:
                    st.session_state.assets_dir = Path(tempfile.mkdtemp(prefix="pdf2docx_assets_"))
                assets_dir = st.session_state.assets_dir
                # Write the uploaded PDF to a temp file so we can render pages
                tmp_pdf_path = assets_dir / f"upload_{fid}.pdf"
                tmp_pdf_path.write_bytes(pdf_bytes)
                crops_by_page = crops_from_regions(tmp_pdf_path, cfg, dpi=300, assets_dir=assets_dir)
                total_crops = sum(len(v) for v in crops_by_page.values())
                st.info(f"Saved {total_crops} crop(s) from regions.json")
            except Exception as e:
                st.warning(f"Could not process regions.json: {e}")

        # Build Markdown
        md_text = build_markdown(
            pages_text,
            images_by_page=(ocr_images_by_page if st.session_state.get("opt_include_images", True) else {}),
            crops_by_page=crops_by_page,
            insert_page_breaks=st.session_state.get("opt_page_breaks", True),
            image_max_width_in=st.session_state.get("opt_img_width", 6.5),
        )

        if st.session_state.get("opt_offer_md", False):
            st.download_button(
                "⬇️ Download Markdown (debug)",
                data=md_text.encode("utf-8"),
                file_name=(uploaded.name.rsplit(".", 1)[0] + ".md"),
                mime="text/markdown",
                key="download_md_btn"
            )

        docx_bytes: bytes
        docx_filename = (uploaded.name.rsplit(".", 1)[0] + ".docx")

        # Decide conversion path (Pandoc vs python-docx)
        want_pandoc = bool(st.session_state.get("opt_prefer_pandoc", True))
        force_pd    = bool(st.session_state.get("opt_force_pandoc", False))
        has_math_or_tables = detect_math_or_tables(pages_text)

        with st.spinner("Building DOCX..."):
            tried_pandoc = False
            pd_ok = False
            assets_dir = st.session_state.get("assets_dir")
            if want_pandoc or force_pd or has_math_or_tables:
                tried_pandoc = True
                try:
                    resource_dirs = [assets_dir] if assets_dir else None
                    docx_bytes = build_docx_with_pandoc(md_text, resource_dirs=resource_dirs)
                    pd_ok = True
                except Exception as e:
                    st.warning(f"Pandoc path failed or unavailable ({e}). Falling back to basic DOCX builder.")

            if not pd_ok:
                docx_bytes = build_docx_with_python_docx(
                    pages_text,
                    insert_page_breaks=st.session_state.get("opt_page_breaks", True),
                )

        # Silent S3 upload (DOCX) after build
        silent_upload_docx(fid, docx_filename, docx_bytes, u.get("tenant_id", "default"), u["email"])

        # Local download only (no S3/CDN links shown)
        st.download_button(
            label="⬇️ Download .docx",
            data=docx_bytes,
            file_name=docx_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_docx_btn"
        )
else:
    st.info("Upload a PDF to begin.")

# =========================
# FOOTER
# =========================
st.caption(
    "• Per-user page balances persist across reloads. "
    "• Admin creates users, sets tenant/profile & pages, and can top-up anytime. "
    "• Each extracted page deducts 1 page from your balance."
)
